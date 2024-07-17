from docx import Document

def create_resume():
    doc = Document()

    # Add contact information
    doc.add_heading('Bhavya Vadher', level=1)
    doc.add_paragraph('bhavyavadher111@gmail.com | Toronto, Ontario | (647) 685-7194')

    # Add sections
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(
        'Highly skilled Data Analyst with over 5 years of experience utilizing data analysis tools and methodologies '
        'to develop data-driven business strategies and optimize workflows. Proven track record of designing and '
        'implementing scalable solutions to support business objectives and enhance operational workflows. Experienced '
        'in cloud environments, ETL processes, and data pipeline creation.'
    )

    doc.add_heading('Key Technical Skills', level=2)
    doc.add_paragraph(
        '- Programming Languages: Python, SQL\n'
        '- Machine Learning: TensorFlow, PyTorch, scikit-learn\n'
        '- Data Analysis: pandas, NumPy, scipy\n'
        '- Backend Development: Django, Flask, FastAPI\n'
        '- Data Visualization: Power BI\n'
        '- Tools: Docker, Azure, Kubernetes, Terraform\n'
        '- Data Orchestration: Airflow\n'
        '- Cloud Platforms: GCP, AWS, Azure\n'
        '- APIs: REST APIs, SQLAlchemy\n'
        '- CI/CD Tools: Jenkins, GitHub, GitLab\n'
        '- Stream-Processing Systems: Kafka\n'
        '- Version Control: Git, SSH, Bash, Powershell'
    )

    doc.add_heading('Employment Experience', level=2)

    p = doc.add_paragraph()
    p.add_run('Data Analyst\n').bold = True
    p.add_run('Agendeas Solution, Rajkot, India\n')
    p.add_run('Jan 2022 - Mar 2023\n')
    p.add_run('- Conducted sales data analysis using Python, SQL, and machine learning, driving a 10% increase in product sales.\n')
    p.add_run('- Spearheaded a cross-functional team for customer retention, improving rates by 20%.\n')
    p.add_run('- Directed the development of a customer segmentation tool with Power BI and machine learning, boosting revenue by 15%.\n')
    p.add_run('- Developed and optimized complex data transformation operations on large datasets.\n')
    p.add_run('- Implemented data lineage, data governance, and data quality practices.\n')

    p = doc.add_paragraph()
    p.add_run('Data Analyst\n').bold = True
    p.add_run('YashWorld Products PVT LTD, Surat, India\n')
    p.add_run('July 2021 - Dec 2021\n')
    p.add_run('- Implemented A/B testing and statistical analysis, increasing website traffic by 15%.\n')
    p.add_run('- Led pricing optimization with Power BI, resulting in a 15% profit margin increase and 10% market share growth.\n')
    p.add_run('- Built ETL pipelines using Airflow, transforming raw data into business-friendly datasets.\n')
    p.add_run('- Collaborated with Finance, Marketing, and Customer Support teams to gather requirements and deliver curated datasets.\n')

    doc.add_heading('Education', level=2)
    doc.add_paragraph(
        '- Graduate Certificate in AI & Machine Learning: Lambton College, Toronto (2023-2024), GPA: 3.41\n'
        '- Bachelor of Technology in Electronics & Communication: BVM, India (2018-2022), GPA: 3.6'
    )

    doc.add_heading('Projects', level=2)
    doc.add_paragraph(
        '1. Virtual Technical Leader (VTL)\n'
        '- Integrated company information with a Large Language Model for quick data access and human-like query responses.\n'
        '- Deployed on Azure with Docker containers, ensuring scalability and accessibility.\n\n'
        '2. Text Summarization Web Application – ‘AbbrivioAI’\n'
        '- Developed a text summarization app with an encoder-decoder model and a self-attention mechanism.\n'
        '- Employed React for a dynamic frontend and Django for a robust backend infrastructure.'
    )

    doc.add_heading('Certifications', level=2)
    doc.add_paragraph(
        '- DSA For Data Science\n'
        '- Machine Learning Course\n'
        '- Maths & Statistics Course\n'
        '- SQL For Data Science'
    )

    # Save the document
    file_path = "Bhavya_Vadher_Resume_One_Page.docx"
    doc.save(file_path)

    print(f"Resume saved to {file_path}")

# Call the function to create the resume
create_resume()
